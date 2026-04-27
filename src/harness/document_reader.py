from __future__ import annotations

import base64
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from harness.llm_client import ImageInput


@dataclass(frozen=True)
class DocumentContent:
    text: str
    images: list[ImageInput] = field(default_factory=list)


def read_document_text(path: str | Path) -> str:
    return read_document_content(path).text


def read_document_content(path: str | Path, *, include_images: bool = False) -> DocumentContent:
    doc_path = Path(path)
    if doc_path.suffix.lower() != ".docx":
        raise ValueError(f"Only .docx files are supported: {doc_path}")
    if not doc_path.exists():
        raise FileNotFoundError(f"Document not found: {doc_path}")

    document = Document(str(doc_path))
    parts: list[str] = []
    images: list[ImageInput] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    for rel in document.part.rels.values():
        target = str(rel.target_ref)
        if any(kind in target for kind in ("image", "chart", "diagram")):
            parts.append(f"[Embedded object detected: {target}]")
        if include_images and hasattr(rel, "target_part"):
            content_type = getattr(rel.target_part, "content_type", "")
            if content_type.startswith("image/"):
                blob = rel.target_part.blob
                images.append(
                    ImageInput(
                        mime_type=content_type,
                        data_base64=base64.b64encode(blob).decode("ascii"),
                    )
                )

    return DocumentContent(text="\n".join(parts).strip(), images=images)


def list_docx_files(directory: str | Path) -> list[Path]:
    base = Path(directory)
    if not base.exists():
        raise FileNotFoundError(f"Directory not found: {base}")
    if not base.is_dir():
        raise NotADirectoryError(f"Expected directory: {base}")
    return sorted(path for path in base.glob("*.docx") if path.is_file() and not path.name.startswith("~$"))
