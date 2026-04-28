from pathlib import Path

from docx import Document
from openpyxl import Workbook

from harness.cli import main
from harness.document_reader import read_document_content, read_document_text
from harness.excel_reader import normalize_scene_key, read_optimization_records
from harness.templates import load_type_mapping, render_scene_template, type_for_file


def test_read_document_text_extracts_paragraphs_and_tables(tmp_path: Path):
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Meeting source paragraph.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Owner"
    table.rows[0].cells[1].text = "Deadline"
    document.save(docx_path)

    text = read_document_text(docx_path)

    assert "Meeting source paragraph." in text
    assert "Owner | Deadline" in text


def test_excel_reader_groups_numbered_scene_labels(tmp_path: Path):
    xlsx_path = tmp_path / "150data.xlsx"
    _write_xlsx(
        xlsx_path,
        [
            ("发布会议1", "标题1", "内容1"),
            ("发布会议2", "标题2", "内容2"),
            ("知识3", "标题3", "内容3"),
        ],
    )

    records = read_optimization_records(xlsx_path)

    assert [record.scene_key for record in records] == ["发布会议", "发布会议", "知识"]
    assert normalize_scene_key("发布会议12") == "发布会议"


def test_type_mapping_uses_keywords_and_file_overrides(tmp_path: Path):
    mapping_file = tmp_path / "类型映射.yaml"
    mapping_file.write_text(
        """
keywords:
  发布会: press_conference
  工作: meeting
files:
  特殊文件.docx: knowledge
""".strip(),
        encoding="utf-8",
    )

    mapping = load_type_mapping(mapping_file)

    assert type_for_file(file_name="新品发布会.docx", file_stem="新品发布会", type_mapping=mapping, default_type="meeting") == "press_conference"
    assert type_for_file(file_name="工作纪要.docx", file_stem="工作纪要", type_mapping=mapping, default_type="news") == "meeting"
    assert type_for_file(file_name="特殊文件.docx", file_stem="特殊文件", type_mapping=mapping, default_type="news") == "knowledge"
    assert type_for_file(file_name="其他.docx", file_stem="其他", type_mapping=mapping, default_type="news") == "news"


def test_scene_template_replaces_requirement_and_format(tmp_path: Path):
    templates_dir = tmp_path / "templates"
    scene_dir = templates_dir / "场景" / "press_conference"
    scene_dir.mkdir(parents=True)
    (templates_dir / "母模板.md").write_text("A\n{requirement}\nB\n{format}", encoding="utf-8")
    (scene_dir / "要求.md").write_text("发布会要求", encoding="utf-8")
    (scene_dir / "格式.md").write_text("发布会格式", encoding="utf-8")

    paths, content = render_scene_template(
        templates_dir=templates_dir,
        base_template_name="母模板.md",
        scene_name="press_conference",
    )

    assert len(paths) == 3
    assert "发布会要求" in content
    assert "发布会格式" in content


def test_optimize_command_outputs_one_template_per_scene_from_xlsx(tmp_path: Path):
    file_processing = tmp_path / "file_processing"
    file_processing.mkdir()
    xlsx_path = file_processing / "150data.xlsx"
    _write_xlsx(
        xlsx_path,
        [
            ("发布会议1", "发布标题1", "The press conference announced a new product."),
            ("发布会议2", "发布标题2", "The press conference answered media questions."),
            ("工作1", "工作标题", "The meeting decided to ship. Alice owns validation by 2026-05-10."),
        ],
    )
    mapping_file = file_processing / "类型映射.yaml"
    mapping_file.write_text("keywords:\n  发布会: press_conference\n  工作: meeting\n", encoding="utf-8")

    templates_dir = tmp_path / "templates"
    (templates_dir / "场景" / "press_conference").mkdir(parents=True)
    (templates_dir / "场景" / "meeting").mkdir(parents=True)
    (templates_dir / "母模板.md").write_text("# 母模板\n\n{requirement}\n\n{format}", encoding="utf-8")
    (templates_dir / "场景" / "press_conference" / "要求.md").write_text("发布会要求", encoding="utf-8")
    (templates_dir / "场景" / "press_conference" / "格式.md").write_text("## 发布内容", encoding="utf-8")
    (templates_dir / "场景" / "meeting" / "要求.md").write_text("工作要求", encoding="utf-8")
    (templates_dir / "场景" / "meeting" / "格式.md").write_text("## 行动项", encoding="utf-8")
    output_dir = tmp_path / "outputs" / "opt"

    code = main(
        [
            "optimize",
            "--file-processing-dir",
            str(file_processing),
            "--optimization-data-file",
            str(xlsx_path),
            "--templates-dir",
            str(templates_dir),
            "--initial-template",
            "母模板.md",
            "--type-mapping-file",
            str(mapping_file),
            "--rules-file",
            "rules/scoring_rules.yaml",
            "--output-dir",
            str(output_dir),
        ]
    )

    assert code == 0
    assert (output_dir / "发布会议" / "final" / "report.md").exists()
    assert (output_dir / "工作" / "final" / "report.md").exists()
    generated = sorted((templates_dir / "generated").glob("*/*/final.md"))
    assert len(generated) == 2
    rendered = "\n".join(path.read_text(encoding="utf-8") for path in generated)
    assert "发布会要求" in rendered
    assert "工作要求" in rendered


def test_optimize_command_auto_creates_initial_template_for_unknown_scene(tmp_path: Path):
    file_processing = tmp_path / "file_processing"
    file_processing.mkdir()
    xlsx_path = file_processing / "150data.xlsx"
    _write_xlsx(xlsx_path, [("新场景1", "标题", "内容")])

    templates_dir = tmp_path / "templates"
    templates_dir.mkdir()
    (templates_dir / "母模板.md").write_text("# 母模板\n\n{requirement}\n\n{format}", encoding="utf-8")
    output_dir = tmp_path / "outputs" / "opt"

    code = main(
        [
            "optimize",
            "--file-processing-dir",
            str(file_processing),
            "--optimization-data-file",
            str(xlsx_path),
            "--templates-dir",
            str(templates_dir),
            "--initial-template",
            "母模板.md",
            "--rules-file",
            "rules/scoring_rules.yaml",
            "--output-dir",
            str(output_dir),
        ]
    )

    assert code == 0
    assert (templates_dir / "场景" / "新场景" / "要求.md").exists()
    assert (templates_dir / "场景" / "新场景" / "格式.md").exists()
    assert list((templates_dir / "generated").glob("*/*/final.md"))


def test_evaluate_command_processes_all_original_docx_with_type_mapping(tmp_path: Path):
    file_processing = tmp_path / "file_processing"
    original_dir = file_processing / "original"
    candidate_a = file_processing / "candidate_a"
    candidate_b = file_processing / "candidate_b"
    original_dir.mkdir(parents=True)
    candidate_a.mkdir()
    candidate_b.mkdir()

    _write_docx(original_dir / "新闻事件.docx", "A news event happened today.")
    _write_docx(original_dir / "知识材料.docx", "This document explains a core concept.")
    for folder in (candidate_a, candidate_b):
        _write_docx(folder / "新闻事件.docx", "A news event summary.")
        _write_docx(folder / "知识材料.docx", "A knowledge summary.")
    mapping_file = file_processing / "类型映射.yaml"
    mapping_file.write_text("keywords:\n  新闻: news\n  知识: knowledge\n", encoding="utf-8")
    output_dir = tmp_path / "outputs" / "eval"

    code = main(
        [
            "evaluate",
            "--file-processing-dir",
            str(file_processing),
            "--type-mapping-file",
            str(mapping_file),
            "--rules-file",
            "rules/scoring_rules.yaml",
            "--output-dir",
            str(output_dir),
        ]
    )

    assert code == 0
    comparison = (output_dir / "comparison.md").read_text(encoding="utf-8")
    assert "news" in comparison
    assert "knowledge" in comparison
    assert len(list((output_dir / "details").glob("*.json"))) == 4


def test_read_document_content_can_extract_images(tmp_path: Path):
    docx_path = tmp_path / "含图文档.docx"
    image_path = tmp_path / "tiny.png"
    image_path.write_bytes(
        bytes.fromhex(
            "89504E470D0A1A0A0000000D4948445200000001000000010802000000907753DE"
            "0000000C4944415408D763F8FFFF3F0005FE02FEA73581E40000000049454E44AE426082"
        )
    )
    document = Document()
    document.add_paragraph("Document with image.")
    document.add_picture(str(image_path))
    document.save(docx_path)

    content = read_document_content(docx_path, include_images=True)

    assert "Document with image." in content.text
    assert len(content.images) == 1
    assert content.images[0].mime_type == "image/png"


def _write_docx(path: Path, text: str) -> None:
    document = Document()
    for line in text.splitlines():
        document.add_paragraph(line)
    document.save(path)


def _write_xlsx(path: Path, rows: list[tuple[str, str, str]]) -> None:
    workbook = Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(row)
    workbook.save(path)
